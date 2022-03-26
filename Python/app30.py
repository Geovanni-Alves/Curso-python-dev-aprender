# Automação de criação de Arquivos Word do zero!
from docx import Document
from docx.shared import Cm

documento = Document()
documento.add_heading("Titulo do documento", 0)

paragrafo = documento.add_paragraph("Um parágrafo simples ")
paragrafo.add_run("e importante").bold = True
paragrafo.add_run(" com palavras em ")
paragrafo.add_run("italico").italic = True

documento.add_heading("Título - Nivel 1", level=1)
documento.add_heading("Título - Nivel 2", level=2)
documento.add_heading("Título - Nivel 3", level=3)

documento.add_paragraph('Formatação "No Spacing"', style="No Spacing")
documento.add_paragraph('Formatação "Heading 1"', style="Heading 1")
documento.add_paragraph('Formatação "Heading 2"', style="Heading 2")
documento.add_paragraph('Formatação "Heading 3"', style="Heading 3")
documento.add_paragraph('Formatação "Title"', style="Title")
documento.add_paragraph('Formatação "Subtitle"', style="Subtitle")
documento.add_paragraph('Formatação "Quote"', style="Quote")
documento.add_paragraph('Formatação "Intense Quote"', style="Intense Quote")
documento.add_paragraph('Formatação "List Paragraph"', style="List Paragraph")
documento.add_paragraph('Formatação "List Bullet"', style="List Bullet")
documento.add_paragraph('Formatação "List Number"', style="List Number")

documento.add_picture("notebook.jpg", width=Cm(5.25))

# tabela = documento.add_table(rows=3, cols=2)
# celula1 = tabela.cell(0, 0)
# celula1.text = "Nome"

compras = (
    (3, "101", "Uva"),
    (7, "422", "Pão"),
    (4, "423", "Banana, Ovos, Tomate, Carne"),
    (5, "433", "Cerveja Becks"),
    (6, "443", "Pipoca")
)

documento.add_page_break()

# Criar a estrutura inicial da tabela
tabela_supermercado = documento.add_table(rows=1, cols=3)
cabecalho_tabela_supermercado = tabela_supermercado.rows[0].cells
cabecalho_tabela_supermercado[0].text = "Quantidade"
cabecalho_tabela_supermercado[1].text = "ID"
cabecalho_tabela_supermercado[2].text = "Descrição"

# Laço para popular a tabela com a quantidade de produtos que tiver
for quantidade, id, desc in compras:
    linha = tabela_supermercado.add_row().cells
    linha[0].text = str(quantidade)
    linha[1].text = id
    linha[2].text = desc


documento.save("demo.docx")
